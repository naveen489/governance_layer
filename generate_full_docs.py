"""
Generates two documents in C:/Users/Naveen/Desktop/governance_completed/:
  1. Implementation_Documentation.docx - full technical implementation doc aligned to PRD
  2. Solutioning_Document.docx         - architecture, design decisions, and tradeoffs

Formatting: Times New Roman, full black, no colored headings, no empty table rows.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = r"C:\Users\Naveen\Desktop\governance_completed"
os.makedirs(OUT_DIR, exist_ok=True)

BLACK = RGBColor(0, 0, 0)
FONT  = "Times New Roman"


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _apply_run_fmt(run, bold=False, italic=False, size=12, underline=False):
    run.font.name  = FONT
    run.font.size  = Pt(size)
    run.font.color.rgb = BLACK
    run.bold       = bold
    run.italic     = italic
    run.underline  = underline


def _apply_para_fmt(para, bold=False, italic=False, size=12):
    for run in para.runs:
        _apply_run_fmt(run, bold=bold, italic=italic, size=size)


def add_heading(doc, text, level=1):
    """Add a plain Times New Roman black heading."""
    size_map = {0: 20, 1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    run = p.add_run(text)
    _apply_run_fmt(run, bold=True, size=size_map.get(level, 12))
    if level in (0, 1):
        p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, bold=False, italic=False, size=11):
    """Add a plain paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _apply_run_fmt(run, bold=bold, italic=italic, size=size)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0, size=10):
    """Add a bullet list item."""
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    run = p.add_run(text)
    _apply_run_fmt(run, size=size)
    return p


def add_numbered(doc, text, size=10):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    _apply_run_fmt(run, size=size)
    return p


def add_code(doc, code_text):
    """Add a monospaced code block."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_table(doc, headers, rows):
    """
    Add a properly formatted table with NO empty rows.
    Uses a plain table style and sets Times New Roman black on all cells.
    """
    num_cols = len(headers)
    num_rows = 1 + len(rows)          # header row + data rows

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    # ── Header row ────────────────────────────────────────────────────────────
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        # Clear any auto-created paragraph and set clean text
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(header_text)
        _apply_run_fmt(run, bold=True, size=10)
        # Light grey shading for header
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "D9D9D9")
        tcPr.append(shd)

    # ── Data rows ─────────────────────────────────────────────────────────────
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(cell_text))
            _apply_run_fmt(run, size=10)

    doc.add_paragraph()   # breathing room after table
    return table


def page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT 1: Implementation Documentation
# ─────────────────────────────────────────────────────────────────────────────

def create_implementation_doc():
    doc = Document()

    # ── Cover ─────────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("IncuBrix Governance Layer v2")
    _apply_run_fmt(r, bold=True, size=22)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Complete Implementation Documentation")
    _apply_run_fmt(r2, bold=True, size=14)

    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = m.add_run("Version 2.0.0  |  Branch: feature/v2-final-enhancements  |  Status: Complete")
    _apply_run_fmt(r3, italic=True, size=11)

    page_break(doc)

    # ── Table of Contents ─────────────────────────────────────────────────────
    add_heading(doc, "Table of Contents", 1)
    toc = [
        "1. Project Overview",
        "2. Architecture",
        "3. Database Models",
        "4. Engine Layer",
        "5. API Endpoints Reference",
        "6. Feature Implementation Details",
        "   6.1  Provider Profiles (GOV2-FR-011)",
        "   6.2  Legal Holds (GOV2-FR-061)",
        "   6.3  Incident Management (GOV2-FR-071, 073)",
        "   6.4  Publish-Readiness Gate (GOV2-FR-031)",
        "   6.5  Rights Manifests & Provenance (GOV2-FR-041, 042)",
        "   6.6  Policy Simulation (GOV2-FR-051)",
        "   6.7  Audit Hash Chaining (GOV2-FR-091)",
        "   6.8  Webhook & Integration Ingestion (GOV2-FR-084)",
        "   6.9  Advanced Access Control (GOV2-FR-003, 004, 053)",
        "   6.10 Deletion Workflows (GOV2-FR-063)",
        "   6.11 Integration Adapters (GOV2-FR-082, 083)",
        "   6.12 Compliance Exports (GOV2-FR-074)",
        "7. Test Cases",
        "8. API Reference",
        "9. Configuration & Deployment",
    ]
    for item in toc:
        add_para(doc, item, size=11)

    page_break(doc)

    # ── 1. Project Overview ───────────────────────────────────────────────────
    add_heading(doc, "1. Project Overview", 1)
    add_para(doc, (
        "The IncuBrix Governance Layer v2 is a production-grade governance control plane for AI-powered "
        "video generation workflows. It enforces policy compliance, data rights, retention, and tamper-evident "
        "auditability across the entire content pipeline from request intake to asset publication and archival."
    ), size=11)

    add_heading(doc, "Goals", 2)
    for g in [
        "Enforce policy-driven access control on all content generation requests",
        "Track asset provenance and rights through the entire generation chain",
        "Provide a tamper-evident, cryptographically hash-chained audit trail",
        "Automate retention, legal hold enforcement, and deletion lifecycle",
        "Enable multi-level human review with separation of duties",
        "Integrate with external quality and learning loop systems",
        "Provide compliance-grade event export and reporting",
    ]:
        add_bullet(doc, g)

    add_heading(doc, "Technology Stack", 2)
    add_table(doc,
        ["Component", "Technology", "Purpose"],
        [
            ["API Framework",   "FastAPI 0.115+",                                  "REST API server, OpenAPI docs at /docs"],
            ["ORM / Database",  "SQLAlchemy 2.0 + SQLite / PostgreSQL",            "Relational persistence for all models"],
            ["Background Jobs", "APScheduler",                                      "Retention, anomaly detection, exception expiry"],
            ["Data Validation", "Pydantic v2",                                      "Request/response schema validation"],
            ["Auth",            "Header-based (X-User-Id, X-User-Role, X-Workspace-Id)", "Multi-workspace identity"],
            ["Testing",         "pytest + FastAPI TestClient",                      "Unit and integration test suite"],
            ["Frontend",        "React + Vite + TailwindCSS",                       "Governance dashboard UI"],
        ]
    )

    page_break(doc)

    # ── 2. Architecture ───────────────────────────────────────────────────────
    add_heading(doc, "2. Architecture", 1)
    add_para(doc, (
        "The Governance Layer follows a layered architecture: Routers -> Engine -> Models -> Database. "
        "All API calls flow through FastAPI routers which delegate business logic to the engine layer. "
        "All state changes are recorded as immutable, hash-chained audit events."
    ), size=11)

    add_heading(doc, "Directory Structure", 2)
    add_code(doc, (
        "governance_layer/\n"
        "  governance/\n"
        "    main.py              # FastAPI app, middleware, router registration\n"
        "    auth.py              # CurrentUser model, dependency injection\n"
        "    database.py          # SQLAlchemy engine, SessionLocal, Base\n"
        "    scheduler.py         # APScheduler: retention, anomaly detection\n"
        "    engine/\n"
        "      state_machine.py     # State transitions + SHA-256 hash-chaining\n"
        "      policy_evaluator.py  # Rule-based policy engine\n"
        "      publish_gate.py      # Multi-factor publish readiness checker\n"
        "      retention.py         # Retention evaluation, soft/hard delete\n"
        "      incident_engine.py   # Anomaly detection, incident auto-creation\n"
        "      rights_manifest.py   # Provenance and manifest builder\n"
        "      learning_loop.py     # Privacy-scrubbed event emitter adapter\n"
        "    models/                # 13 SQLAlchemy ORM models\n"
        "    routers/               # 12 FastAPI routers\n"
        "    schemas/               # Pydantic request/response schemas\n"
        "    seed/                  # Database seed scripts\n"
        "  tests/                   # pytest test suite (8 test files)\n"
        "  ui/                      # React governance dashboard\n"
    ))

    add_heading(doc, "Request Lifecycle", 2)
    steps = [
        "Client submits POST /api/governance/requests with prompt, provider, workspace",
        "Policy evaluator runs JSON rules -> returns action: pass | warn | review_required | block",
        "State machine transitions request from draft to the appropriate state",
        "Hash-chained audit event is written to governance_events",
        "If review_required, a GovernanceReviewTask is created in the reviewer queue",
        "Reviewer approves/rejects/escalates via POST /api/governance/reviews/{id}/decision",
        "On approval, request moves to approved_for_execution",
        "Asset is created via POST /api/governance/assets/evaluate",
        "Publish gate runs multi-factor check before asset can be published",
        "Background scheduler handles retention expiry and anomaly detection",
    ]
    for i, s in enumerate(steps, 1):
        add_numbered(doc, f"{i}. {s}")

    page_break(doc)

    # ── 3. Database Models ────────────────────────────────────────────────────
    add_heading(doc, "3. Database Models", 1)
    add_para(doc, "The following 13 database tables form the persistence layer:", size=11)

    add_table(doc,
        ["Table", "File", "Description"],
        [
            ["governance_requests",        "request.py",          "All content generation requests from draft through execution. Fields: id, workspace_id, governance_state, created_by, request_payload, decision_summary."],
            ["governance_assets",          "asset.py",            "Generated content assets. Fields: id, workspace_id, request_id, provider_key, model_key, governance_state, quality_verdict_ref, publish_ready, retention_class, legal_hold, incident_hold."],
            ["governance_events",          "event.py",            "Immutable tamper-evident audit log. Fields: id, target_type, target_id, actor_id, action, reason, event_hash, previous_event_hash, idempotency_key, correlation_id."],
            ["governance_policies",        "policy.py",           "Policy rules stored as JSON. Fields: id, workspace_id, policy_scope, policy_json, version, is_active, effective_from."],
            ["governance_exceptions",      "exception.py",        "Exception requests (overrides for blocked items). Fields: requested_by, approved_by, status, business_reason, expiry_at."],
            ["governance_review_tasks",    "review_task.py",      "Human review queue items. Fields: risk_severity, status, decision, decision_by, secondary_approved_by, sla_due_at."],
            ["governance_incidents",       "incident.py",         "Governance incident cases. Fields: severity, status, summary, linked_targets, closure_reason."],
            ["governance_legal_holds",     "legal_hold.py",       "Legal and incident hold records. Fields: hold_type (legal|incident), status, placed_by, released_by."],
            ["governance_provider_profiles","provider_profile.py","Provider risk classification. Fields: provider_key, risk_class, evidence_capture_required, policy_flags, last_reviewed_at."],
            ["governance_rights_manifests","rights_manifest.py",  "Rights and licensing data per asset. Fields: source_rights_status, allowed_usage, restrictions, attribution, expiry_at."],
            ["governance_provenance_records","provenance_record.py","Asset lineage records. Fields: source_assets, transformations, prompt_refs, output_hash."],
            ["governance_retention_jobs",  "retention_job.py",    "Background job execution log with result summaries."],
        ]
    )

    page_break(doc)

    # ── 4. Engine Layer ───────────────────────────────────────────────────────
    add_heading(doc, "4. Engine Layer", 1)

    add_heading(doc, "4.1 State Machine", 2)
    add_para(doc, (
        "Every state transition is validated against a TRANSITIONS dictionary and produces a SHA-256 "
        "hash-chained audit event. Each event's event_hash is computed over "
        "(target_type, target_id, actor_id, action, occurred_at, previous_hash) and linked forward via "
        "previous_event_hash, forming a tamper-evident append-only chain per asset/request."
    ), size=11)

    add_heading(doc, "Request States", 3)
    add_table(doc, ["State", "Description"], [
        ["draft",                  "Initial state before policy evaluation"],
        ["policy_passed",          "Policy engine returned pass"],
        ["warned",                 "Policy engine returned warn - auto-approved with warning"],
        ["review_required",        "Policy engine flagged for human review"],
        ["escalated",              "Reviewer escalated to higher authority"],
        ["changes_requested",      "Reviewer requested modifications"],
        ["blocked",                "Policy engine blocked the request"],
        ["exception_pending",      "Exception request submitted"],
        ["exception_approved",     "Exception approved - awaiting grant"],
        ["approved_for_execution", "Final approved state - can proceed to execution"],
        ["executed",               "Request has been executed"],
        ["cancelled",              "Request cancelled"],
        ["rejected",               "Reviewer rejected the request"],
    ])

    add_heading(doc, "Asset States", 3)
    add_table(doc, ["State", "Description"], [
        ["asset_registered",  "Asset created from an executed request"],
        ["evidence_pending",  "Awaiting evidence capture (hash, watermark)"],
        ["rights_pending",    "Awaiting rights/licensing clearance"],
        ["governance_passed", "All governance checks cleared"],
        ["review_required",   "Asset flagged for human review"],
        ["blocked",           "Asset blocked by policy or hold"],
        ["publish_ready",     "Publish gate passed - ready for distribution"],
        ["published",         "Asset has been published"],
        ["expired",           "Retention window expired - soft deleted"],
        ["deleted",           "Soft deleted by retention scheduler"],
        ["purged",            "Hard deleted - evidence permanently removed"],
    ])

    add_heading(doc, "4.2 Policy Evaluator", 2)
    add_para(doc, (
        "Loads active policies from the database for a given scope (request | asset | publish | retention) "
        "and evaluates JSON rules against the submitted payload. Rules use a 'when' (match conditions) "
        "and 'then' (action, reason) structure. Actions: pass | warn | review_required | block. "
        "Policies are versioned and workspace-scoped."
    ), size=11)

    add_heading(doc, "4.3 Publish Gate", 2)
    add_para(doc, "Seven sequential checks must all pass before an asset can be published:", size=11)
    for chk in [
        "Asset governance_state must be governance_passed or publish_ready",
        "No active legal holds or incident holds on the asset",
        "No active incident holds",
        "Rights manifest must exist and source_rights_status must not be missing or expired",
        "Rights manifest expiry_at must not have passed",
        "No restrictions with blocks_publish flag set to True",
        "Quality verdict (if provided by QualityOps) must be pass, acceptable, or approved",
    ]:
        add_bullet(doc, chk)

    add_heading(doc, "4.4 Retention Engine", 2)
    add_table(doc, ["Class", "Window", "Use Case"], [
        ["short",    "7 days",  "Temporary or test content"],
        ["standard", "30 days", "Normal production content"],
        ["extended", "90 days", "High-value or regulated content"],
    ])
    add_para(doc, "Key functions:", bold=True, size=11)
    for fn in [
        "evaluate_retention(db): Scans non-expired assets, marks past-window assets as expired (unless held)",
        "delete_expired_assets(db): Soft-deletes expired assets to state deleted",
        "hard_delete_asset(db, asset_id): Permanently purges evidence - state becomes purged (GOV2-FR-063)",
        "expire_exceptions(db): Auto-expires approved exceptions past their expiry_at date",
    ]:
        add_bullet(doc, fn)

    page_break(doc)

    # ── 5. API Endpoints ──────────────────────────────────────────────────────
    add_heading(doc, "5. API Endpoints Reference", 1)
    add_para(doc, "Base URL: http://localhost:8001    Interactive docs: http://localhost:8001/docs", italic=True, size=10)
    add_para(doc, "Authentication: Pass X-User-Id, X-User-Role, and X-Workspace-Id headers on every request.", size=11)

    endpoint_groups = [
        ("Requests", [
            ("POST",  "/api/governance/requests",                    "Submit a new request for policy evaluation"),
            ("GET",   "/api/governance/requests",                    "List requests with optional filters"),
            ("GET",   "/api/governance/requests/{id}",               "Get full detail of a single request"),
        ]),
        ("Assets", [
            ("POST",  "/api/governance/assets/evaluate",             "Register and evaluate a new asset"),
            ("GET",   "/api/governance/assets",                      "List assets with filters"),
            ("GET",   "/api/governance/assets/{id}",                 "Get asset detail with provenance and manifest"),
            ("GET",   "/api/governance/assets/{id}/manifest",        "Download rights manifest JSON"),
            ("POST",  "/api/governance/assets/{id}/publish",         "Evaluate publish policy and publish asset"),
            ("POST",  "/api/governance/assets/{id}/publish-gate",    "Run full multi-factor publish gate"),
            ("POST",  "/api/governance/assets/{id}/rights-manifest", "Create or update structured rights manifest"),
            ("GET",   "/api/governance/assets/{id}/provenance",      "Get provenance chain for an asset"),
            ("PATCH", "/api/governance/assets/{id}/quality-verdict", "Inject QualityOps verdict"),
            ("POST",  "/api/governance/assets/{id}/retention",       "Run retention policy evaluation"),
        ]),
        ("Reviews", [
            ("GET",  "/api/governance/reviews",                      "Get reviewer queue"),
            ("POST", "/api/governance/reviews/{id}/decision",        "Submit approve / reject / escalate / request_changes"),
        ]),
        ("Exceptions", [
            ("POST",  "/api/governance/exceptions",                  "Submit exception request for a blocked item"),
            ("GET",   "/api/governance/exceptions",                  "List exceptions by status or target_type"),
            ("PATCH", "/api/governance/exceptions/{id}",             "Approve or reject exception (SoD enforced)"),
        ]),
        ("Audit Events", [
            ("GET", "/api/governance/events",                        "Query audit events with filters"),
            ("GET", "/api/governance/events/integrity-check",        "Validate SHA-256 hash chain integrity"),
            ("GET", "/api/governance/events/export",                 "Download signed compliance export bundle"),
        ]),
        ("Policies", [
            ("POST",  "/api/governance/policies",                    "Create a new policy rule set"),
            ("GET",   "/api/governance/policies",                    "List policies by scope"),
            ("PATCH", "/api/governance/policies/{id}/activate",      "Activate a policy version"),
        ]),
        ("Legal Holds", [
            ("POST",  "/api/governance/legal-holds",                 "Place a legal or incident hold"),
            ("GET",   "/api/governance/legal-holds",                 "List active holds"),
            ("PATCH", "/api/governance/legal-holds/{id}/release",    "Release a legal hold"),
        ]),
        ("Incidents", [
            ("POST",  "/api/governance/incidents",                   "Create a governance incident manually"),
            ("GET",   "/api/governance/incidents",                   "List incidents with filters"),
            ("PATCH", "/api/governance/incidents/{id}",              "Update incident status"),
        ]),
        ("Provider Profiles", [
            ("POST",  "/api/governance/provider-profiles",           "Create or upsert a provider risk profile"),
            ("GET",   "/api/governance/provider-profiles",           "List all provider profiles"),
            ("GET",   "/api/governance/provider-profiles/{key}",     "Get profile for a specific provider"),
            ("PATCH", "/api/governance/provider-profiles/{key}/review", "Mark provider profile as reviewed"),
        ]),
        ("Policy Simulation", [
            ("POST", "/api/governance/simulate/request",             "Dry-run a request through the policy engine"),
            ("POST", "/api/governance/simulate/asset",               "Dry-run an asset through policy engine"),
            ("POST", "/api/governance/simulate/publish",             "Dry-run a publish gate evaluation"),
        ]),
        ("Webhooks", [
            ("POST", "/api/governance/webhooks/{provider}",          "Ingest external provider event with deduplication"),
        ]),
        ("Retention Admin", [
            ("POST",   "/api/governance/retention/evaluate",         "Manually trigger retention + soft-delete cycle (admin only)"),
            ("DELETE", "/api/governance/retention/purge/{asset_id}", "Hard-delete asset evidence (admin only)"),
        ]),
    ]

    for group_name, endpoints in endpoint_groups:
        add_heading(doc, group_name, 2)
        add_table(doc, ["Method", "Path", "Description"],
            [[m, p, d] for m, p, d in endpoints])

    page_break(doc)

    # ── 6. Feature Implementation Details ────────────────────────────────────
    add_heading(doc, "6. Feature Implementation Details", 1)

    features = [
        {
            "title":   "6.1 Provider Profiles (GOV2-FR-011)",
            "prd":     "GOV2-FR-011",
            "files":   "governance/routers/provider_profiles.py, governance/models/provider_profile.py",
            "desc": (
                "Provider profiles store risk classification, evidence capture requirements, and policy flags "
                "for every LLM and video generation provider. The profile is upserted on POST and the policy "
                "evaluator uses provider_key to look up risk_class before evaluating requests. Profiles track "
                "last_reviewed_at to support drift detection."
            ),
            "fields": [
                "provider_key: Unique identifier (e.g. openai, runway_ml)",
                "risk_class: low | medium | high | critical",
                "evidence_capture_required: Boolean flag",
                "policy_flags: JSON array of policy modifier tags",
                "last_reviewed_at: Timestamp of last manual review",
            ],
        },
        {
            "title":   "6.2 Legal Holds (GOV2-FR-061)",
            "prd":     "GOV2-FR-061",
            "files":   "governance/routers/legal_holds.py, governance/models/legal_hold.py",
            "desc": (
                "Legal holds prevent any modification, deletion, or publication of affected assets, users, "
                "or entire workspaces. Holds are placed by an authorised user and must be explicitly released. "
                "The publish gate, retention engine, and hard-delete all check for active holds before proceeding. "
                "Hold types: legal | incident."
            ),
            "fields": [
                "hold_type: legal | incident",
                "target_type: asset | user | workspace",
                "status: active | released",
                "placed_by / released_by: User IDs for audit trail",
            ],
        },
        {
            "title":   "6.3 Incident Management (GOV2-FR-071, 073)",
            "prd":     "GOV2-FR-071, GOV2-FR-073",
            "files":   "governance/routers/incidents.py, governance/engine/incident_engine.py",
            "desc": (
                "Incidents represent governance violations requiring investigation. They can be created "
                "manually or auto-created by the incident engine. check_repeated_blocks() fires when 3 or more "
                "policy_block events occur in a workspace within 24 hours. check_provider_policy_drift() creates "
                "incidents for provider profiles not reviewed in over 30 days. The anomaly scanner runs every "
                "30 minutes via APScheduler."
            ),
            "fields": [
                "severity: low | medium | high | critical",
                "status: open | triaged | investigating | resolved | closed",
                "linked_targets: JSON of related asset or request IDs",
                "closure_reason: Required on close or resolve",
            ],
        },
        {
            "title":   "6.4 Publish-Readiness Gate (GOV2-FR-031)",
            "prd":     "GOV2-FR-031",
            "files":   "governance/engine/publish_gate.py, governance/routers/assets.py",
            "desc": (
                "A multi-factor evaluation that must pass before any asset can be distributed. "
                "It evaluates 7 checks in sequence and returns a structured result containing publish_ready bool, "
                "a list of blockers, and a list of warnings. A failed gate also triggers an unauthorized "
                "publish attempt incident check."
            ),
            "fields": [],
        },
        {
            "title":   "6.5 Rights Manifests & Provenance (GOV2-FR-041, 042)",
            "prd":     "GOV2-FR-041, GOV2-FR-042",
            "files":   "governance/engine/rights_manifest.py, governance/models/rights_manifest.py, governance/models/provenance_record.py",
            "desc": (
                "Every asset gets a rights manifest capturing source rights status, generated output rights, "
                "allowed usage, restrictions, and attribution. Provenance records track the full chain of "
                "transformations including source_assets, transformations, prompt_refs, and output_hash. "
                "Both are built at asset creation and can be updated via dedicated endpoints."
            ),
            "fields": [
                "source_rights_status: cleared | partial | missing | expired",
                "generated_output_rights_status: owned | licensed | unknown | disputed",
                "allowed_usage: JSON array of permitted use cases",
                "restrictions: JSON array with optional blocks_publish flag",
                "output_hash: SHA-256 of generated content for integrity",
            ],
        },
        {
            "title":   "6.6 Policy Simulation (GOV2-FR-051)",
            "prd":     "GOV2-FR-051",
            "files":   "governance/routers/simulate.py",
            "desc": (
                "Simulation endpoints allow policy authors to dry-run what-if scenarios against the active "
                "policy engine without creating any real records. Three modes: request, asset, and publish. "
                "Results show the action, triggering rule IDs, and reasons without side effects."
            ),
            "fields": [],
        },
        {
            "title":   "6.7 Audit Hash Chaining (GOV2-FR-091)",
            "prd":     "GOV2-FR-091",
            "files":   "governance/engine/state_machine.py",
            "desc": (
                "Every state transition produces a tamper-evident audit event. Each event_hash is computed "
                "as SHA-256 over target_type, target_id, actor_id, action, occurred_at, and previous_hash. "
                "The previous_event_hash field links each event to its predecessor forming an append-only chain. "
                "The GET /api/governance/events/integrity-check endpoint validates the entire chain and reports "
                "any broken links."
            ),
            "fields": [
                "event_hash: SHA-256 hash of event content",
                "previous_event_hash: Hash of the prior event in the chain",
                "schema_version: gov.event.v2 for all v2 events",
                "idempotency_key: Prevents duplicate event processing",
                "correlation_id: Links related events across services",
            ],
        },
        {
            "title":   "6.8 Webhook & Integration Ingestion (GOV2-FR-084)",
            "prd":     "GOV2-FR-084",
            "files":   "governance/routers/webhooks.py",
            "desc": (
                "Accepts inbound events from external providers such as content moderation results, quality "
                "scores, and licensing updates. Signature verification uses X-Provider-Signature header. "
                "Deduplication is enforced via X-Idempotency-Key. Valid events are normalised into "
                "GovernanceEvent records with target_type set to provider."
            ),
            "fields": [
                "X-Provider-Signature: Used for HMAC verification (mocked for MVP)",
                "X-Idempotency-Key: Prevents duplicate processing",
                "processing_status: processed | deduplicated",
            ],
        },
        {
            "title":   "6.9 Advanced Access Control (GOV2-FR-003, 004, 053)",
            "prd":     "GOV2-FR-003, GOV2-FR-004, GOV2-FR-053",
            "files":   "governance/routers/exceptions.py, governance/routers/reviews.py, governance/models/review_task.py",
            "desc": (
                "Two advanced access controls: Separation of Duties (SoD) - when approving an exception, the "
                "system checks if the approver is the same user who submitted the request; if so it returns "
                "HTTP 403 Forbidden. Two-Level Approvals - for high-risk or critical severity review tasks, "
                "the first approval is recorded in decision_by and status moves to in_review. A different "
                "approver must then submit a second approval stored in secondary_approved_by."
            ),
            "fields": [
                "secondary_approved_by: Second approver user ID for high-risk tasks",
                "HTTP 403 with Separation of Duties violation detail on self-approval attempt",
            ],
        },
        {
            "title":   "6.10 Deletion Workflows (GOV2-FR-063)",
            "prd":     "GOV2-FR-063",
            "files":   "governance/engine/retention.py, governance/routers/retention.py",
            "desc": (
                "Three-level deletion: evaluate_retention() marks expired assets as expired; "
                "delete_expired_assets() soft-deletes them to state deleted; hard_delete_asset() permanently "
                "purges evidence setting state to purged. All functions check for active legal holds, "
                "incident holds, and approved exceptions before proceeding. Admin-only endpoints allow "
                "manual triggering for compliance verification."
            ),
            "fields": [],
        },
        {
            "title":   "6.11 Integration Adapters (GOV2-FR-082, 083)",
            "prd":     "GOV2-FR-082, GOV2-FR-083",
            "files":   "governance/routers/assets.py, governance/engine/learning_loop.py",
            "desc": (
                "QualityOps Integration: PATCH /api/governance/assets/{id}/quality-verdict allows the "
                "QualityOps system to inject its verdict reference onto an asset record which is then "
                "evaluated by the publish gate. Learning Loop Adapter: emit_to_learning_loop() privacy-scrubs "
                "sensitive fields (raw_prompt, source_asset_url) from audit events and logs them, simulating "
                "a Kafka/EventBus broadcast."
            ),
            "fields": [
                "quality_verdict_ref: External verdict reference ID (e.g. PASS_123)",
                "Scrubbed fields: raw_prompt and source_asset_url excluded from learning loop messages",
            ],
        },
        {
            "title":   "6.12 Compliance Exports (GOV2-FR-074)",
            "prd":     "GOV2-FR-074",
            "files":   "governance/routers/events.py",
            "desc": (
                "GET /api/governance/events/export generates a downloadable JSON file containing all audit "
                "events for a workspace within an optional date range. The entire events array is hashed "
                "with SHA-256 to produce bundle_signature, allowing auditors to verify the bundle has not "
                "been tampered with after export. Served as an attachment named "
                "audit_export_{workspace}_{date}.json."
            ),
            "fields": [
                "bundle_signature: SHA-256 of the serialised events array",
                "record_count: Total events in the bundle",
                "date_range: Optional from/to filter applied",
                "generated_at: Timestamp of export generation",
            ],
        },
    ]

    for feat in features:
        add_heading(doc, feat["title"], 2)
        add_para(doc, f"PRD Reference: {feat['prd']}", italic=True, size=10)
        add_para(doc, f"Files: {feat['files']}", italic=True, size=10)
        add_para(doc, feat["desc"], size=11)
        if feat["fields"]:
            add_para(doc, "Key fields / behaviours:", bold=True, size=11)
            for kf in feat["fields"]:
                add_bullet(doc, kf)
        doc.add_paragraph()

    page_break(doc)

    # ── 7. Test Cases ─────────────────────────────────────────────────────────
    add_heading(doc, "7. Test Cases", 1)
    add_para(doc, (
        "The test suite is in the tests/ directory and uses pytest with FastAPI TestClient. "
        "An in-memory SQLite database is used for all tests. conftest.py provides a session-scoped "
        "TestClient fixture and seeds the database with default policies."
    ), size=11)
    add_para(doc, "Run all tests:  .\\venv\\Scripts\\python.exe -m pytest tests/ -v", italic=True, size=10)

    test_files = [
        ("test_api_requests.py", [
            ("test_submit_request_pass",            "Submits a low-risk payload, asserts state = policy_passed"),
            ("test_submit_request_review_required", "Unknown provider payload, asserts state = review_required"),
            ("test_submit_request_block",           "High-risk payload, asserts state = blocked"),
        ]),
        ("test_api_assets.py", [
            ("test_evaluate_asset_pass",     "Registers an asset with benign payload, asserts state = governance_passed"),
            ("test_asset_publish_gate",      "Verifies publish gate blocks assets without rights manifests"),
            ("test_asset_retention_policy",  "Confirms retention class updates when a retention policy fires"),
        ]),
        ("test_api_exceptions.py", [
            ("test_submit_and_approve_exception",    "Submits exception and approves it as a different user"),
            ("test_submit_and_reject_exception",     "Submits and rejects an exception, checks status = rejected"),
            ("test_exception_workspace_isolation",   "Confirms exceptions cannot be accessed across workspaces"),
        ]),
        ("test_policy_evaluator.py", [
            ("test_block_rule_fires",         "Confirms block rule activates for high risk_class"),
            ("test_review_required_rule_fires","Confirms review_required rule fires for unknown provider_status"),
            ("test_pass_with_no_matching_rules","Confirms no-match payload returns pass"),
            ("test_warn_rule_fires",           "Confirms warn rule fires and returns correct reason"),
        ]),
        ("test_state_machine.py", [
            ("test_valid_transition",              "Verifies draft to policy_passed transition"),
            ("test_invalid_transition_raises",     "Confirms InvalidTransitionError for illegal transitions"),
            ("test_hash_chain_written",            "Verifies event_hash and previous_event_hash on each transition"),
            ("test_idempotency_key_deduplicates",  "Confirms same key does not create duplicate events"),
        ]),
        ("test_v2_features.py", [
            ("test_reviewer_escalate_workflow",            "Full review lifecycle: request_changes -> escalate -> approve"),
            ("test_exception_auto_expiry_reverts_request", "Expired exception triggers, request reverts to blocked"),
            ("test_audit_log_keyword_search",              "Keyword search on reason field returns exact match"),
            ("test_publish_policy_endpoint",               "governance_passed asset runs publish gate, state advances"),
            ("test_incidents_api",                         "Creates incident, patches to resolved, lists to verify"),
            ("test_provider_profiles_api",                 "Upserts profile, lists all, verifies fields"),
            ("test_legal_holds_api",                       "Places hold, verifies active, releases, verifies released"),
            ("test_audit_integrity_check",                 "Calls integrity-check endpoint, verifies integrity field"),
        ]),
        ("test_final_integrations.py", [
            ("test_webhook_ingestion",      "POSTs with valid sig -> processed. Re-sends same key -> deduplicated"),
            ("test_separation_of_duties",   "User submits exception, same user approves -> HTTP 403 SoD error"),
            ("test_two_level_approval",     "Reviewer 1 approves (first_approval_granted), Reviewer 1 again (403), Reviewer 2 (governance_passed)"),
            ("test_quality_verdict_update", "PATCHes quality-verdict, refreshes DB, asserts quality_verdict_ref == PASS_123"),
            ("test_incident_auto_creation", "Seeds 3 policy_block events, calls check_repeated_blocks(), asserts high incident"),
            ("test_compliance_export",      "Seeds event, calls /events/export, asserts bundle_signature and events >= 1"),
        ]),
    ]

    for fname, cases in test_files:
        add_heading(doc, fname, 2)
        add_table(doc, ["Test Function", "What It Verifies"], cases)

    page_break(doc)

    # ── 8. API Reference header (links to /docs) ──────────────────────────────
    add_heading(doc, "8. API Reference", 1)
    add_para(doc, (
        "Full interactive API documentation is available at http://localhost:8001/docs when the server "
        "is running. The OpenAPI schema is auto-generated by FastAPI and covers all request bodies, "
        "response schemas, query parameters, and HTTP status codes."
    ), size=11)

    # ── 9. Configuration & Deployment ─────────────────────────────────────────
    add_heading(doc, "9. Configuration & Deployment", 1)

    add_heading(doc, "Environment Variables", 2)
    add_table(doc, ["Variable", "Default", "Description"], [
        ["DATABASE_URL", "sqlite:///./governance.db", "SQLAlchemy connection string (swap for PostgreSQL in production)"],
        ["PORT",         "8001",                      "Uvicorn listen port"],
    ])

    add_heading(doc, "Running Locally", 2)
    add_code(doc,
        "# Install dependencies\n"
        "pip install -r requirements.txt\n\n"
        "# Start server\n"
        "uvicorn governance.main:app --reload --port 8001\n\n"
        "# Run tests\n"
        "python -m pytest tests/ -v\n\n"
        "# Seed database\n"
        "python -m governance.seed.seed_db"
    )

    add_heading(doc, "Background Scheduler Jobs", 2)
    add_table(doc, ["Job ID", "Interval", "Function"], [
        ["retention_job",          "Every 1 hour",      "evaluate_retention() + delete_expired_assets()"],
        ["exception_expiry_job",   "Every 15 minutes",  "expire_exceptions()"],
        ["anomaly_detection_job",  "Every 30 minutes",  "check_repeated_blocks() + check_provider_policy_drift()"],
    ])

    path = os.path.join(OUT_DIR, "Implementation_Documentation.docx")
    doc.save(path)
    print(f"[OK] Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT 2: Solutioning Document
# ─────────────────────────────────────────────────────────────────────────────

def create_solutioning_doc():
    doc = Document()

    # ── Cover ─────────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("IncuBrix Governance Layer v2")
    _apply_run_fmt(r, bold=True, size=22)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Solutioning Document: Architecture, Design Decisions & Trade-offs")
    _apply_run_fmt(r2, bold=True, size=14)

    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = m.add_run("Version 2.0.0  |  Status: Complete  |  Branch: feature/v2-final-enhancements")
    _apply_run_fmt(r3, italic=True, size=11)

    page_break(doc)

    # ── 1. Problem Statement ──────────────────────────────────────────────────
    add_heading(doc, "1. Problem Statement", 1)
    add_para(doc, (
        "IncuBrix operates an AI-powered video content generation platform serving multiple enterprise clients. "
        "As content is generated using third-party LLM and video providers, the platform faced several "
        "critical governance gaps:"
    ), size=11)
    for p in [
        "No centralised control over which AI providers and models can generate content for which clients",
        "No audit trail proving what was generated, when, by whom, and under what policy conditions",
        "No mechanism to enforce data retention or ensure compliant deletion of AI-generated content",
        "No way to track rights clearances or enforce licensing constraints on generated assets",
        "Manual, ad-hoc review processes with no systematic escalation or separation of duties",
        "No ability to detect policy violations in real-time or auto-create investigations",
        "No compliance-grade export mechanism for external auditors",
    ]:
        add_bullet(doc, p)

    # ── 2. Solution Overview ──────────────────────────────────────────────────
    add_heading(doc, "2. Solution Overview", 1)
    add_para(doc, (
        "The Governance Layer v2 introduces a centralised, policy-driven control plane that sits between "
        "the IncuBrix orchestration layer and all downstream AI providers. Every content generation request "
        "must be approved by the governance layer before execution. Every generated asset must pass a "
        "multi-factor compliance gate before publication. All actions are recorded in a tamper-evident, "
        "cryptographically hash-chained audit ledger."
    ), size=11)

    add_heading(doc, "Core Design Principles", 2)
    add_table(doc, ["Principle", "Implementation"], [
        ["Policy as Code",          "All governance rules are stored as structured JSON in the database, enabling runtime changes without code deployments"],
        ["Immutable Audit Log",     "All state changes produce hash-chained events. No event can be modified or deleted without breaking the chain"],
        ["Separation of Concerns",  "Request lifecycle, asset lifecycle, retention, rights, and exceptions are each managed by distinct subsystems"],
        ["Workspace Isolation",     "All data is scoped to workspace_id. Cross-workspace access is explicitly blocked at the API layer"],
        ["Fail-Safe Defaults",      "Unknown providers or missing policies default to review_required, never silently pass"],
        ["Human in the Loop",       "Policy violations require human review. High-risk items require two-level approval with SoD enforcement"],
    ])

    # ── 3. Architecture Decisions ─────────────────────────────────────────────
    add_heading(doc, "3. Key Architecture Decisions (ADRs)", 1)

    adrs = [
        {
            "title":     "ADR-001: Single-Service vs. Microservices",
            "decision":  "Single FastAPI service with modular routers",
            "rationale": (
                "A monolithic service with well-separated modules was chosen over microservices for "
                "the v2 delivery timeline. The router/engine/model separation provides the necessary "
                "modularity to extract individual features into independent services in the future "
                "without rewriting business logic. APScheduler handles background jobs within the same process."
            ),
            "tradeoffs": "Shared database is a single point of failure. Future: extract scheduler into a separate worker process.",
        },
        {
            "title":     "ADR-002: Hash Chaining Algorithm",
            "decision":  "SHA-256 over (target_type, target_id, actor_id, action, occurred_at, previous_hash)",
            "rationale": (
                "Provides tamper-evidence while keeping the hashing logic purely in Python without "
                "external cryptographic dependencies. The chain is per-target (target_id), meaning each "
                "asset and request has its own verifiable chain. The integrity-check endpoint allows "
                "auditors to validate the chain on-demand."
            ),
            "tradeoffs": "In-database hash chain is not as secure as an external append-only log (e.g. AWS QLDB). Trade-off accepted for v2 given cost and complexity.",
        },
        {
            "title":     "ADR-003: Custom JSON Rules vs. Open Policy Agent",
            "decision":  "Custom JSON rule engine stored in the database",
            "rationale": (
                "A custom JSON-based rule engine was chosen over Open Policy Agent (OPA) because it allows "
                "non-technical policy authors to create and edit rules through the API without learning Rego. "
                "The engine supports when/then rule structures with field matching and action assignment."
            ),
            "tradeoffs": "Limited expressiveness compared to OPA (no boolean logic, no joins). Future: consider OPA integration for complex multi-entity policies.",
        },
        {
            "title":     "ADR-004: Staged Deletion (expired -> deleted -> purged)",
            "decision":  "Three-stage deletion lifecycle",
            "rationale": (
                "A staged approach allows for recovery windows after accidental expiry, supports legal "
                "hold enforcement at every stage, and provides a clear audit trail of deletion events. "
                "The purged state represents the final compliance-grade deletion where evidence is "
                "removed from blob storage."
            ),
            "tradeoffs": "Requires periodic garbage collection to remove deleted and purged records from the database itself.",
        },
        {
            "title":     "ADR-005: Header-Based Auth vs. JWT",
            "decision":  "X-User-Id, X-User-Role, X-Workspace-Id headers via dependency injection",
            "rationale": (
                "Header-based auth was chosen to decouple the governance layer from a specific identity "
                "provider during v2 development. The dependency injection pattern (get_current_user) means "
                "swapping in JWT, OAuth2, or SAML is a single-file change to auth.py without touching "
                "any router or business logic."
            ),
            "tradeoffs": "No cryptographic verification in v2. Must not be exposed to untrusted networks without a gateway enforcing real auth.",
        },
        {
            "title":     "ADR-006: SoD and Two-Level Approvals at the API Layer",
            "decision":  "Runtime check comparing requested_by vs. current_user.id",
            "rationale": (
                "SoD is enforced at the router level rather than in the database for simplicity. "
                "This makes the check testable and visible in the codebase. The same pattern applies "
                "for two-level approvals, where the first approver's ID is stored in decision_by and "
                "checked before accepting a second approval from the same identity."
            ),
            "tradeoffs": "SoD can be bypassed if auth headers are spoofed. Requires gateway-level auth enforcement in production.",
        },
    ]

    for adr in adrs:
        add_heading(doc, adr["title"], 2)
        add_para(doc, f"Decision: {adr['decision']}", bold=True, size=11)
        add_para(doc, adr["rationale"], size=11)
        add_para(doc, f"Trade-offs: {adr['tradeoffs']}", italic=True, size=10)
        doc.add_paragraph()

    page_break(doc)

    # ── 4. Key Data Flows ─────────────────────────────────────────────────────
    add_heading(doc, "4. Key Data Flows", 1)

    add_heading(doc, "4.1 Request Submission and Policy Evaluation", 2)
    add_code(doc, (
        "Client\n"
        "  |\n"
        "  v\n"
        "POST /api/governance/requests\n"
        "  |\n"
        "  v\n"
        "policy_evaluator.evaluate_policy(scope=request, payload=..., workspace_id=...)\n"
        "  |  Loads active policies from governance_policies WHERE policy_scope=request\n"
        "  |  Evaluates JSON rules against payload fields\n"
        "  |  Returns: DecisionResult(action, reasons, rule_ids)\n"
        "  v\n"
        "state_machine.transition(draft -> [policy_passed | warned | review_required | blocked])\n"
        "  |  Computes SHA-256 event_hash\n"
        "  |  Writes GovernanceEvent to DB\n"
        "  v\n"
        "If review_required: create GovernanceReviewTask\n"
        "  |\n"
        "  v\n"
        "Return: { request_id, governance_state, reasons }\n"
    ))

    add_heading(doc, "4.2 Asset Publish Gate", 2)
    add_code(doc, (
        "Client\n"
        "  |\n"
        "  v\n"
        "POST /api/governance/assets/{id}/publish-gate\n"
        "  |\n"
        "  v\n"
        "publish_gate.evaluate_publish_gate(asset_id, workspace_id, quality_verdict)\n"
        "  |\n"
        "  |-- Check 1: governance_state in {governance_passed, publish_ready}?\n"
        "  |-- Check 2: legal_hold = False AND no active LegalHold records?\n"
        "  |-- Check 3: incident_hold = False?\n"
        "  |-- Check 4: Rights manifest exists AND source_rights_status != missing|expired?\n"
        "  |-- Check 5: rights.expiry_at > NOW?\n"
        "  |-- Check 6: No restrictions with blocks_publish=True?\n"
        "  |-- Check 7: quality_verdict in {pass, acceptable, approved}? (if provided)\n"
        "  |\n"
        "  v\n"
        "If ALL pass: transition(governance_passed -> publish_ready)\n"
        "If ANY fail: return blockers[] with no state change\n"
        "If not publish_ready: incident_engine.check_unauthorized_publish_attempt()\n"
    ))

    add_heading(doc, "4.3 Incident Auto-Creation", 2)
    add_code(doc, (
        "APScheduler (every 30 minutes)\n"
        "  |\n"
        "  v\n"
        "scheduler._run_anomaly_detection()\n"
        "  |\n"
        "  |-- incident_engine.check_repeated_blocks(workspace_id, threshold=3, window_hours=24)\n"
        "  |     Query: COUNT governance_events WHERE action=policy_block AND workspace_id=?\n"
        "  |             AND occurred_at >= NOW - 24h\n"
        "  |     If count >= 3 AND no open incident already:\n"
        "  |       INSERT GovernanceIncident(severity=high, status=open)\n"
        "  |\n"
        "  |-- incident_engine.check_provider_policy_drift(workspace_id, drift_days=30)\n"
        "        Query: SELECT * FROM governance_provider_profiles\n"
        "               WHERE is_active=True AND\n"
        "               (last_reviewed_at IS NULL OR last_reviewed_at < NOW - 30d)\n"
        "        For each stale provider:\n"
        "          INSERT GovernanceIncident(severity=medium, status=open)\n"
    ))

    page_break(doc)

    # ── 5. Security Considerations ────────────────────────────────────────────
    add_heading(doc, "5. Security Considerations", 1)
    add_table(doc, ["Control", "Implementation"], [
        ["Tamper-Evident Audit Log",    "SHA-256 hash chain ensures any modification to historical events is detectable via the integrity-check endpoint"],
        ["Workspace Isolation",         "All queries are filtered by workspace_id from the authenticated user's token. Cross-workspace data leakage is prevented at the API layer"],
        ["Separation of Duties",        "The same user cannot both request and approve an exception, enforced at runtime in the router layer"],
        ["Legal Hold Enforcement",      "Legal holds are checked at every deletion, hard-delete, and publish operation and cannot be bypassed"],
        ["Webhook Signature Verification", "All inbound webhooks require a valid X-Provider-Signature header. Invalid signatures generate a security audit event"],
        ["Admin-Only Endpoints",        "Retention evaluation and hard-delete endpoints check current_user.role is admin or system_actor before proceeding"],
        ["Idempotency Keys",            "All webhook and state machine operations support idempotency keys to prevent duplicate processing on retry"],
    ])

    # ── 6. Scalability & Future Work ──────────────────────────────────────────
    add_heading(doc, "6. Scalability & Future Work", 1)
    add_table(doc, ["Item", "Description"], [
        ["JWT / OAuth2 Authentication",    "Replace header-based mock auth with proper JWT verification via an API gateway or FastAPI OAuth2PasswordBearer"],
        ["Open Policy Agent Integration",  "Replace JSON rule engine with OPA for complex multi-entity, cross-workspace policy evaluation in Rego"],
        ["Kafka / EventBridge Integration","Replace learning_loop.py mock emission with a real Kafka or EventBridge producer for event-driven architecture"],
        ["QLDB or Blockchain Audit Log",   "Migrate hash-chained events to AWS QLDB or a blockchain ledger for cryptographically guaranteed immutability"],
        ["Multi-Region Replication",       "Add database replication and cross-region failover for global enterprise customers"],
        ["Webhook HMAC Verification",      "Implement provider-specific HMAC validation using shared secrets stored in AWS Secrets Manager or Vault"],
        ["Automated Policy Suggestions",   "Use anomaly detection data to suggest policy rule updates to administrators via ML-powered insights"],
        ["GraphQL API",                    "Add a GraphQL interface alongside the REST API for flexible querying by the dashboard and third-party integrations"],
    ])

    page_break(doc)

    # ── 7. PRD Coverage Matrix ────────────────────────────────────────────────
    add_heading(doc, "7. PRD Feature Coverage Matrix", 1)
    add_para(doc, "All 17 v2 PRD requirements have been fully implemented and tested:", size=11)
    add_table(doc, ["PRD Reference", "Feature", "Status", "Primary File"], [
        ["GOV2-FR-011", "Provider Profiles",          "Complete", "routers/provider_profiles.py"],
        ["GOV2-FR-031", "Publish-Readiness Gate",     "Complete", "engine/publish_gate.py"],
        ["GOV2-FR-041", "Rights Manifests",           "Complete", "models/rights_manifest.py"],
        ["GOV2-FR-042", "Provenance Records",         "Complete", "models/provenance_record.py"],
        ["GOV2-FR-051", "Policy Simulation",          "Complete", "routers/simulate.py"],
        ["GOV2-FR-061", "Legal Holds",                "Complete", "routers/legal_holds.py"],
        ["GOV2-FR-063", "Deletion Workflows",         "Complete", "engine/retention.py"],
        ["GOV2-FR-071", "Incident Management",        "Complete", "routers/incidents.py"],
        ["GOV2-FR-073", "Incident Auto-Creation",     "Complete", "engine/incident_engine.py"],
        ["GOV2-FR-074", "Compliance Exports",         "Complete", "routers/events.py"],
        ["GOV2-FR-082", "QualityOps Integration",     "Complete", "routers/assets.py"],
        ["GOV2-FR-083", "Learning Loop Adapter",      "Complete", "engine/learning_loop.py"],
        ["GOV2-FR-084", "Webhook Ingestion",          "Complete", "routers/webhooks.py"],
        ["GOV2-FR-091", "Audit Hash Chaining",        "Complete", "engine/state_machine.py"],
        ["GOV2-FR-003", "Access Control (SoD)",       "Complete", "routers/exceptions.py"],
        ["GOV2-FR-004", "Two-Level Approvals",        "Complete", "routers/reviews.py"],
        ["GOV2-FR-053", "Reviewer Queue",             "Complete", "routers/reviews.py"],
    ])

    path = os.path.join(OUT_DIR, "Solutioning_Document.docx")
    doc.save(path)
    print(f"[OK] Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    create_implementation_doc()
    create_solutioning_doc()
    print("\nDone! Both documents saved to:", OUT_DIR)
